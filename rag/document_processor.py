"""
Procesador de documentos (PDF, DOCX, TXT, CSV, XLSX)
"""
import os
import io
from typing import List, Dict, Tuple
import logging
from pathlib import Path

# Document parsers
import PyPDF2
from docx import Document
import pandas as pd

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Procesador unificado de documentos"""
    
    @staticmethod
    def extract_text(file_content: bytes, filename: str) -> str:
        """Extraer texto según tipo de archivo"""
        ext = Path(filename).suffix.lower()
        
        try:
            if ext == '.pdf':
                return DocumentProcessor._extract_pdf(file_content)
            elif ext == '.docx':
                return DocumentProcessor._extract_docx(file_content)
            elif ext == '.txt':
                return file_content.decode('utf-8')
            elif ext == '.csv':
                return DocumentProcessor._extract_csv(file_content)
            elif ext == '.xlsx':
                return DocumentProcessor._extract_xlsx(file_content)
            elif ext == '.md':
                return file_content.decode('utf-8')
            else:
                raise ValueError(f"Tipo de archivo no soportado: {ext}")
                
        except Exception as e:
            logger.error(f"Error extrayendo texto de {filename}: {e}")
            raise
    
    @staticmethod
    def _extract_pdf(file_content: bytes) -> str:
        """Extraer texto de PDF"""
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text_parts = []
        for page_num, page in enumerate(pdf_reader.pages, 1):
            text = page.extract_text()
            if text.strip():
                text_parts.append(f"--- Página {page_num} ---\n{text}")
        
        return "\n\n".join(text_parts)
    
    @staticmethod
    def _extract_docx(file_content: bytes) -> str:
        """Extraer texto de DOCX"""
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extraer tablas
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        return "\n\n".join(text_parts)
    
    @staticmethod
    def _extract_csv(file_content: bytes) -> str:
        """Extraer texto de CSV"""
        csv_file = io.BytesIO(file_content)
        df = pd.read_csv(csv_file)
        
        # Convertir a texto estructurado
        text_parts = [
            f"CSV con {len(df)} filas y {len(df.columns)} columnas",
            f"Columnas: {', '.join(df.columns.tolist())}",
            "\nDatos:",
            df.to_string(index=False, max_rows=100)
        ]
        
        return "\n".join(text_parts)
    
    @staticmethod
    def _extract_xlsx(file_content: bytes) -> str:
        """Extraer texto de XLSX"""
        xlsx_file = io.BytesIO(file_content)
        
        text_parts = []
        sheets = pd.read_excel(xlsx_file, sheet_name=None)
        
        for sheet_name, df in sheets.items():
            text_parts.append(f"\n=== Hoja: {sheet_name} ===")
            text_parts.append(f"{len(df)} filas, {len(df.columns)} columnas")
            text_parts.append(f"Columnas: {', '.join(df.columns.tolist())}")
            text_parts.append(df.to_string(index=False, max_rows=50))
        
        return "\n".join(text_parts)
    
    @staticmethod
    def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
        """Dividir texto en chunks con overlap"""
        if not text.strip():
            return []
        
        # Dividir por párrafos primero
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        chunks = []
        current_chunk = []
        current_size = 0
        
        for para in paragraphs:
            para_size = len(para)
            
            if current_size + para_size <= chunk_size:
                current_chunk.append(para)
                current_size += para_size
            else:
                # Guardar chunk actual
                if current_chunk:
                    chunks.append("\n\n".join(current_chunk))
                
                # Iniciar nuevo chunk con overlap
                if overlap > 0 and current_chunk:
                    overlap_text = current_chunk[-1][-overlap:]
                    current_chunk = [overlap_text, para]
                    current_size = len(overlap_text) + para_size
                else:
                    current_chunk = [para]
                    current_size = para_size
        
        # Agregar último chunk
        if current_chunk:
            chunks.append("\n\n".join(current_chunk))
        
        return chunks
